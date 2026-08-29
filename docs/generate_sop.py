#!/usr/bin/env python3
"""Generate IDP SOP Word document with embedded screenshots."""

from __future__ import annotations

import json
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SHOTS = DOCS / "sop-screenshots"
OUT = DOCS / "SOP-IDP-Maintenance_Extraction.docx"
BASE = "http://127.0.0.1:8000"
API = "http://127.0.0.1:8001"


def login_payload() -> tuple[str, dict]:
    req = urllib.request.Request(
        f"{API}/api/auth/login",
        data=json.dumps({"email": "admin@omniparse.local", "password": "ChangeMeNow!"}).encode(),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=10) as resp:
        data = json.loads(resp.read().decode())
    return data["access_token"], data["user"]


def capture_admin_screenshots() -> None:
    token, user = login_payload()
    user_json = json.dumps(user)
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(viewport={"width": 1920, "height": 1080})
        page = context.new_page()

        def seed_auth(url: str) -> None:
            page.goto(url, wait_until="networkidle")
            page.evaluate(
                """([token, user]) => {
                    localStorage.setItem('omniparse_auth_token', token);
                    localStorage.setItem('omniparse_auth_user', user);
                }""",
                [token, user_json],
            )
            page.reload(wait_until="networkidle")
            page.wait_for_timeout(1200)
            # Close login/landing overlay if shown
            close_btn = page.locator("#login-close, #landing-close, button:has-text('×')")
            if close_btn.count():
                try:
                    close_btn.first.click(timeout=2000)
                    page.wait_for_timeout(400)
                except Exception:
                    pass

        seed_auth(f"{BASE}/")
        page.click("#rail-chat-btn", timeout=10000)
        page.wait_for_timeout(600)
        page.screenshot(path=str(SHOTS / "fig07-ai-assistant.png"))

        seed_auth(f"{BASE}/admin.html")
        page.wait_for_selector("#admin-monitor-grid", timeout=15000)
        page.wait_for_timeout(1500)
        page.screenshot(path=str(SHOTS / "fig08-admin-monitoring.png"), full_page=True)

        page.click('button[data-admin-view="users"]')
        page.wait_for_timeout(800)
        page.screenshot(path=str(SHOTS / "fig09-admin-users.png"), full_page=True)

        page.click('button[data-admin-view="logs"]')
        page.wait_for_timeout(800)
        page.screenshot(path=str(SHOTS / "fig10-admin-extraction-logs.png"), full_page=True)

        seed_auth(f"{BASE}/history.html")
        page.wait_for_selector("#history-summary", timeout=15000)
        page.wait_for_timeout(1000)
        page.screenshot(path=str(SHOTS / "fig11-my-extracts.png"), full_page=True)

        browser.close()


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph("STANDARD OPERATING PROCEDURE")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    title = doc.add_paragraph("IDP Maintenance Manual Extraction and Review")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(20)

    sub = doc.add_paragraph(
        "O&M manual and logbook intake, AI extraction, registry review, export and admin oversight"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta = [
        ("Document ID", "SOP-IDP-001"),
        ("Version", "1.0"),
        ("Prepared Date", "11 August 2026"),
        ("Effective Date", "Upon approval"),
        ("Process Owner", "Engineering / Maintenance Data Operations"),
        ("Classification", "Internal Use Only"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    for i, (k, v) in enumerate(meta):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()
    p = doc.add_paragraph("CONTROLLED DOCUMENT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    note = doc.add_paragraph(
        "Printed copies are uncontrolled unless formally stamped and issued."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_figure(doc: Document, num: int, caption: str, image: str) -> None:
    path = SHOTS / image
    doc.add_paragraph(f"Figure {num} — {caption}")
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
    else:
        doc.add_paragraph(f"[Screenshot placeholder: {image}]")
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()


def build_document() -> None:
    doc = Document()
    add_cover(doc)

    add_heading(doc, "1. Document Control")
    add_heading(doc, "1.1 Approval", 2)
    add_table(
        doc,
        ["Role", "Name", "Signature", "Date"],
        [
            ["Prepared by – Process/Automation Team", "", "", ""],
            ["Reviewed by – Maintenance Data Owner", "", "", ""],
            ["Approved by – Engineering Management", "", "", ""],
        ],
    )

    add_heading(doc, "1.2 Revision History", 2)
    add_table(
        doc,
        ["Version", "Date", "Author", "Description of Change", "Status"],
        [
            [
                "1.0",
                "11 Aug 2026",
                "Maintenance Data Operations / Automation Team",
                "Initial standard SOP for IDP maintenance manual extraction.",
                "Draft for Approval",
            ]
        ],
    )

    add_heading(doc, "1.3 Document Governance", 2)
    for line in [
        "Review Frequency: At least annually, or immediately after a process, control, configuration or system change.",
        "Controlled Repository: Approved corporate document repository / Engineering Operations knowledge base.",
        "Record Retention: In accordance with the applicable maintenance data and audit-retention policy.",
        "Related Procedures: SharePoint/Fabric sync (future phase); AWS deployment runbook; incident management procedure.",
        "System Versions: IDP UI (CloudFront/S3) + FastAPI API (ECS Fargate); Gemini API extraction engine.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "2. Contents")
    for item in [
        "1. Document Control",
        "2. Contents",
        "3. Purpose and Objectives",
        "4. Scope",
        "5. Roles and Responsibilities",
        "6. Preconditions, Access and Required Inputs",
        "7. Definitions and System Components",
        "8. Process Overview and Decision Rules",
        "9. Detailed Operating Procedure",
        "10. Validation and Control Checklist",
        "11. Exception Handling and Escalation",
        "12. Records, Evidence and Completion Criteria",
    ]:
        doc.add_paragraph(item)

    add_heading(doc, "3. Purpose and Objectives")
    for line in [
        "This SOP defines the standard method for uploading engineering manuals into IDP, running AI extraction, reviewing maintenance/spare-parts/troubleshooting registries, and exporting approved data to Excel.",
        "Provide a consistent and auditable extraction and review process for O&M manuals and field history cards / logbooks.",
        "Ensure extracted data is validated against the source manual before export.",
        "Prevent unsupported edits, duplicate exports, and incomplete exception handling.",
        "Define the minimum evidence required to demonstrate successful processing or re-run.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "4. Scope")
    doc.add_paragraph("This procedure applies to:")
    for line in [
        "Case 1 — O&M manuals processed with Gemini API (production path).",
        "Case 2 — Field history cards / logbooks (Logbook document type).",
        "Case 3 — Failed or low-quality extractions requiring re-run or escalation.",
        "Case 4 — Global Admin monitoring, user management, and extraction log review.",
        "Case 5 — My extracts history review and audit traceability.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(
        "Out of Scope: SharePoint → Fabric automated sync (future phase); AWS infrastructure deployment; ERP/master-data load of exported Excel; procurement or payment processes."
    )

    add_heading(doc, "5. Roles and Responsibilities")
    add_table(
        doc,
        ["Role", "Responsibility"],
        [
            ["Extractor / Reviewer", "Uploads manuals, validates registries against source, edits permitted fields, exports Excel."],
            ["Maintenance SME", "Confirms technical accuracy of tasks, intervals, part numbers, and troubleshooting steps."],
            ["Global Admin", "Creates users, assigns models and Copilot limits, monitors ops and extraction logs."],
            ["Platform / Automation Support", "Investigates API, ECS, Gemini, and integration errors without bypassing controls."],
        ],
    )

    add_heading(doc, "6. Preconditions, Access and Required Inputs")
    for line in [
        "Valid IDP user account (production: https://d11bl7hg497hj.cloudfront.net).",
        "Supported browser (Chrome / Edge / Firefox) with stable connectivity.",
        "Header status shows Python API Ready for production extractions.",
        "Gemini API key configured server-side (Secrets Manager) or as permitted by Global Admin.",
        "Source manual complete and legible; PDF preferred. Word documents are converted to PDF by the API.",
        "Optional page range available for large manuals (CloudFront read timeout max 180 seconds).",
        "Control Requirement: Never amend extracted values merely to force completion. Corrections must be supported by the manual image or approved SME confirmation.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "7. Definitions and System Components")
    terms = [
        ("IDP", "Intelligent Document Processing application for O&M manuals and history cards."),
        ("Python API", "FastAPI backend that performs PDF/OCR processing and Gemini extraction."),
        ("Gemini API", "Cloud LLM engine used for structured registry extraction."),
        ("OCR Vision", "Parse strategy for scanned manuals and history cards."),
        ("Native parse", "Parse strategy for digital/text-based PDFs."),
        ("Registry", "Editable tables: Maintenance, Spare parts, Troubleshooting."),
        ("AI Assistant (Copilot)", "Q&A panel for questions about the loaded manual."),
        ("My extracts", "User history of prior extraction runs."),
        ("Global Admin", "Admin console for users, monitoring, and extraction logs."),
        ("Extract audit", "Server-side audit trail stored in S3 and visible in Admin → Extraction logs."),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        p.add_run(term).bold = True
        p.add_run(f" — {definition}")

    add_heading(doc, "8. Process Overview and Decision Rules")
    doc.add_paragraph(
        "The operator signs in, configures the parsing engine, uploads a manual, reviews registries, and exports Excel only when validation passes."
    )
    add_heading(doc, "8.1 Decision Matrix", 2)
    add_table(
        doc,
        ["Condition", "Route", "Required Action", "Completion Evidence"],
        [
            ["Standard O&M PDF", "O&M + Gemini + Native", "Upload full manual or page range; review all tabs.", "Excel export + My extracts status = pass."],
            ["Scanned manual", "O&M + Gemini + OCR Vision", "Use OCR; split into page ranges if timeout risk.", "Same as above."],
            ["Logbook / history cards", "Logbook document type", "Prefer OCR Vision; review logbook stats.", "Export + My extracts pass."],
            ["API unreachable", "Heuristics fallback only", "Do not use for production; escalate.", "Screenshot + incident ticket."],
            ["Quality below threshold", "Any", "Re-run or obtain SME sign-off; do not export.", "Extraction log entry + notes."],
        ],
    )

    add_heading(doc, "9. Detailed Operating Procedure")
    add_heading(doc, "9.1 Access and Configure IDP", 2)
    steps_91 = [
        (1, "Open IDP", "Navigate to the production URL (or local http://localhost:8000 for pilot).", "fig01-main-workspace.png", "IDP main workspace."),
        (2, "Sign in", "Select Sign in. Enter email and password provided by Global Admin. Select Sign in once.", "fig02-sign-in-modal.png", "Sign-in screen."),
        (3, "Confirm API status", "Verify the header/engine area indicates Python API Ready before production extraction.", "fig01-main-workspace.png", "Header and workspace."),
        (4, "Open AI Parsing Engine Settings", "Select the settings control. Choose O&M Manual or Field History / Logbook. Select the assigned Gemini model.", "fig03-ai-parsing-settings.png", "AI Parsing Engine Settings panel."),
    ]
    for num, title, text, img, cap in steps_91:
        doc.add_paragraph(f"{num}\t{title} — {text}")
        add_figure(doc, num, cap, img)

    add_heading(doc, "9.2 Case 1 — O&M Manual Extraction (Production Path)", 2)
    steps_92 = [
        (5, "Open Upload", "Select Upload. Optionally enter From page and To page for large manuals.", "fig04-upload-page-range.png", "Upload document with optional page range."),
        (6, "Choose file", "Select Choose file and pick PDF (preferred), Word, TXT, or image. Wait for the progress overlay to complete. Do not close the browser tab.", "fig01-main-workspace.png", "Workspace during/after upload."),
        (7, "Review Maintenance Tasks", "Select Maintenance Tasks. Compare Equipment Title, routine, checks, and page references to the manual.", "fig01-main-workspace.png", "Maintenance registry tab."),
        (8, "Review Spare Parts", "Select Spare Parts & Components. Use type filters (Critical, Consumable, Standard). Validate part numbers and quantities.", "fig05-spare-parts-tab.png", "Spare parts registry with filters."),
        (9, "Review Troubleshooting", "Select Troubleshooting. Confirm problem, root-cause/solution, and page references.", "fig06-troubleshooting-tab.png", "Troubleshooting registry tab."),
        (10, "Use AI Assistant (optional)", "Open AI Assistant to query the loaded manual within the daily Copilot limit.", "fig07-ai-assistant.png", "AI Assistant panel."),
        (11, "Export to Excel", "When all validations pass, select Export Excel (3 sheets). Save the file with a traceable name.", "fig01-main-workspace.png", "Export control on registry toolbar."),
        (12, "Confirm in My extracts", "Open profile menu → My extracts. Confirm the run shows pass status.", "fig11-my-extracts.png", "My extracts history page."),
    ]
    for num, title, text, img, cap in steps_92:
        doc.add_paragraph(f"{num}\t{title} — {text}")
        add_figure(doc, num, cap, img)

    add_heading(doc, "9.3 Case 2 — Logbook / History Cards", 2)
    doc.add_paragraph(
        "13\tSelect Logbook type — In AI Parsing Engine Settings, choose Field History / Logbook. Prefer OCR Vision for scanned cards."
    )
    doc.add_paragraph(
        "14\tUpload and review — Follow the same upload, review, and export steps as Case 1, validating logbook-specific columns and stats."
    )
    add_figure(doc, 13, "Logbook document type selection", "fig03-ai-parsing-settings.png")

    add_heading(doc, "9.4 Case 3 — Re-run, Cancel, or Reject", 2)
    doc.add_paragraph(
        "15\tCancel stuck jobs — If progress does not advance, select Cancel on the progress overlay. Do not refresh repeatedly."
    )
    doc.add_paragraph(
        "16\tAdjust and re-run — Change page range, parse strategy (Native vs OCR Vision), or model only when supported by the manual type."
    )
    doc.add_paragraph(
        "17\tDo not export unchecked data — If registries are empty or quality is unacceptable, do not export. Record the issue and escalate if needed."
    )
    doc.add_paragraph(
        "Re-run Quality: Do not export and re-upload the same manual without documenting why the first run failed."
    )

    add_heading(doc, "9.5 Global Admin — Monitoring, Users, and Logs", 2)
    steps_95 = [
        (18, "Open Admin console", "Sign in as Global Admin. Open Admin console from the profile menu.", "fig08-admin-monitoring.png", "Admin Monitoring overview."),
        (19, "Review Monitoring", "Check Overview, App ops, and AWS services cards. Use Refresh to update metrics.", "fig08-admin-monitoring.png", "Monitoring dashboard."),
        (20, "Manage users", "Select Users. Create users, assign allowed models, set Copilot daily limits, enable or disable accounts.", "fig09-admin-users.png", "Users directory."),
        (21, "Review extraction logs", "Select Extraction logs. Investigate failed runs, low quality scores, and duration anomalies.", "fig10-admin-extraction-logs.png", "Extraction logs table."),
    ]
    for num, title, text, img, cap in steps_95:
        doc.add_paragraph(f"{num}\t{title} — {text}")
        add_figure(doc, num, cap, img)

    doc.add_paragraph(
        "Screenshot Note: Values shown in figures are examples from the pilot environment. Users must process live manuals and live account settings in production."
    )

    add_heading(doc, "9.6 Concurrent Uploads — How the Queue Works", 2)
    doc.add_paragraph(
        "IDP processes one extraction at a time to avoid overloading the API and Gemini rate limits. "
        "When multiple users upload documents simultaneously, jobs are automatically queued in the order they were submitted. "
        "No upload is rejected — each user's job waits its turn and completes when the slot is free."
    )
    doc.add_paragraph(
        "What the user sees in the browser:"
    )
    for line in [
        'The progress overlay appears immediately after upload with the message: "Python API job queued… Queued — waiting for previous extraction to finish".',
        "The progress bar stays at the queued position until the job ahead finishes.",
        "Once the slot is free the job starts automatically — the user does not need to do anything.",
        "The job completes and registries load exactly as for a single user.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_figure(doc, 13, "User view — job queued in browser while another extraction is running.", "fig13-queue-ui-user-view.png")
    add_figure(doc, 14, "API proof — 3 users accepted simultaneously (HTTP 200, queue positions 1–3), all completed in order.", "fig12-queue-proof.png")

    doc.add_paragraph(
        "Control: Do not cancel and re-upload while a job is queued. The cancel button stops your own job but does not skip the queue for a new one. "
        "Wait for the queued job to complete or cancel it and re-upload only if the settings need to change."
    )

    add_heading(doc, "10. Validation and Control Checklist")
    add_table(
        doc,
        ["No.", "Control", "Required Check", "Done"],
        [
            ["1", "Document identity", "Filename and upload refer to the correct manual.", "[ ]"],
            ["2", "Document type", "O&M Manual vs Logbook selected correctly.", "[ ]"],
            ["3", "API status", "Python API Ready (production).", "[ ]"],
            ["4", "Engine / strategy", "Gemini model and Native/OCR appropriate for source.", "[ ]"],
            ["5", "Page range", "Covers intended section when used.", "[ ]"],
            ["6", "Maintenance", "Tasks, intervals, and checks supported by manual.", "[ ]"],
            ["7", "Spare parts", "Part numbers, types, and quantities plausible.", "[ ]"],
            ["8", "Troubleshooting", "Problems and solutions complete and sourced.", "[ ]"],
            ["9", "Row counts", "Reasonable for manual size; no obvious gaps.", "[ ]"],
            ["10", "Quality", "Quality score acceptable or SME sign-off recorded.", "[ ]"],
            ["11", "Export", "Excel opens with populated sheets.", "[ ]"],
            ["12", "Traceability", "My extracts / audit log shows pass or documented error.", "[ ]"],
        ],
    )

    add_heading(doc, "11. Exception Handling and Escalation")
    add_table(
        doc,
        ["Exception", "Immediate Action", "Escalate / Evidence", "Priority"],
        [
            ["Python API not reachable", "Check /api/health; retry once; do not rely on Heuristics for production.", "Platform Support; attach health check result.", "High"],
            ["CloudFront timeout (180s)", "Split manual into smaller page ranges and re-run.", "Platform Support if persistent.", "High"],
            ["Gemini key invalid", "Verify Secrets Manager or admin-assigned key.", "Global Admin.", "High"],
            ["Empty registries", "Try OCR Vision; verify page range and file quality.", "Maintenance SME.", "Medium"],
            ["Stuck progress overlay", "Cancel job; check Admin extract queue.", "Platform Support.", "Medium"],
            ["Low quality score", "SME review; re-extract with adjusted settings.", "Process Owner.", "Medium"],
            ["Copilot limit reached", "Wait for daily reset or request limit increase.", "Global Admin.", "Low"],
        ],
    )
    add_heading(doc, "11.1 Minimum Incident Information", 2)
    for line in [
        "User email, date/time, filename, engine, parse strategy, page range.",
        "Exact error message, screenshot, and My extracts run reference if available.",
        "Action taken immediately before the error.",
        "Whether a partial export or duplicate run may exist.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(
        "Critical Control: A technical error does not automatically mean extraction failed. Check My extracts and Admin logs before re-uploading."
    )

    add_heading(doc, "12. Records, Evidence and Completion Criteria")
    doc.add_paragraph("Minimum traceable records:")
    for line in [
        "Extract audit entry: user, filename, status, row counts, duration, quality score.",
        "Exported Excel file with traceable naming convention.",
        "SME review notes when a quality exception is approved.",
        "Incident/ticket reference for unresolved errors.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "12.1 Completion Criteria", 2)
    for label, text in [
        ("Complete", "Registries reviewed, Excel exported and saved, My extracts shows pass."),
        ("Incomplete", "Error or abandoned run logged; escalated if production-blocking."),
        ("Escalated", "Not exported; evidence attached; owner and next action recorded."),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)

    doc.add_paragraph(
        "End of Procedure: The transaction is complete only when final status is verified and required evidence is retained."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    import sys
    SHOTS.mkdir(parents=True, exist_ok=True)
    if "--doc-only" not in sys.argv:
        capture_admin_screenshots()
    build_document()
